from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from datetime import date

ROOT = Path(__file__).resolve().parent
IMG = ROOT / 'guide-images'
OUT = ROOT / 'THISO_Leasing_Platform_Huong_Dan_Su_Dung.docx'
BLUE = '163B65'; LIGHT = 'E8EEF5'; NAVY = RGBColor(22, 59, 101); MUTED = RGBColor(90, 104, 120)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for name, size, before, after, color in [('Title', 28, 0, 10, NAVY), ('Heading 1', 16, 18, 10, NAVY), ('Heading 2', 13, 14, 7, NAVY), ('Heading 3', 12, 10, 5, RGBColor(31,77,120))]:
    st = styles[name]; st.font.name = 'Calibri'; st.font.size = Pt(size); st.font.color.rgb = color; st.font.bold = True
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]
header.text = 'THISO LEASING PLATFORM  |  HƯỚNG DẪN SỬ DỤNG'
header.runs[0].font.size = Pt(8); header.runs[0].font.color.rgb = MUTED
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
field = OxmlElement('w:fldSimple'); field.set(qn('w:instr'), 'PAGE'); footer._p.append(field)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet'); p.add_run(text); p.paragraph_format.space_after = Pt(4); return p

def add_step(n, title, detail):
    p = doc.add_paragraph(); p.paragraph_format.keep_with_next = True
    r = p.add_run(f'Bước {n}. {title}'); r.bold = True; r.font.color.rgb = NAVY
    p2 = doc.add_paragraph(detail); p2.paragraph_format.left_indent = Inches(.18); p2.paragraph_format.space_after = Pt(8)

def add_image(filename, caption):
    path = IMG / filename
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(5.7))
    c = doc.add_paragraph(caption); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True; c.runs[0].font.size = Pt(9); c.runs[0].font.color.rgb = MUTED
    c.paragraph_format.keep_with_next = False

def add_module(title, image, intro, steps, notes=None):
    heading = doc.add_heading(title, level=1); heading.paragraph_format.page_break_before = True; doc.add_paragraph(intro)
    add_image(image, f'Hình: Giao diện {title}')
    doc.add_heading('Cách thao tác', level=2)
    for i, (t, d) in enumerate(steps, 1): add_step(i, t, d)
    if notes:
        doc.add_heading('Lưu ý', level=2)
        for note in notes: add_bullet(note)

# Cover — editorial_cover pattern
doc.add_paragraph('SỔ TAY VẬN HÀNH ERP', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(110)
r = p.add_run('THISO Leasing Platform'); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY
p = doc.add_paragraph('Hướng dẫn sử dụng chi tiết theo quy trình end-to-end'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(15); p.runs[0].font.color.rgb = MUTED
p = doc.add_paragraph('\nDành cho Ban điều hành, Leasing, Vận hành, Tài chính, Quản trị hệ thống và Khách thuê'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(45)
p = doc.add_paragraph(f'Phiên bản UAT • {date.today().strftime("%d/%m/%Y")}'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True; p.runs[0].font.color.rgb = NAVY

doc.add_page_break(); doc.add_heading('Mục lục nhanh', level=1)
for line in ['1. Đăng nhập và nguyên tắc sử dụng', '2. Quy trình ERP end-to-end', '3. Dashboard và Mặt bằng', '4. CRM, Booking, Proposal, Approval, Contract', '5. Fitout và Ticket', '6. Billing, Announcements và Tenant Portal', '7. Quản trị hệ thống', '8. Checklist vận hành và xử lý sự cố', '9. Phụ lục kết quả QC']:
    add_bullet(line)
doc.add_heading('Đối tượng và quyền', level=2)
for role, scope in [('Super Admin', 'Toàn quyền hệ thống, tài khoản, Mall, phân quyền và cấu hình.'), ('Leasing/Manager', 'CRM, booking, đề xuất, phê duyệt, hợp đồng.'), ('Operation', 'Fitout, ticket, bảo trì, thông báo vận hành.'), ('Finance', 'Hóa đơn, công nợ, đối soát, SAP.'), ('Tenant', 'Hợp đồng, hóa đơn, yêu cầu hỗ trợ, Fitout và thông báo của Mall đang thuê.')]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    lead = p.add_run(f'{role}: '); lead.bold = True; lead.font.color.rgb = NAVY
    p.add_run(scope)

doc.add_page_break(); doc.add_heading('1. Đăng nhập và nguyên tắc sử dụng', level=1)
add_step(1, 'Mở hệ thống', 'Truy cập địa chỉ hệ thống do quản trị viên cung cấp. UAT hiện dùng http://localhost:58080.')
add_step(2, 'Nhập tài khoản', 'Nhập email và mật khẩu được cấp, sau đó chọn Đăng nhập. Không chia sẻ tài khoản giữa nhiều người.')
add_step(3, 'Chọn Mall', 'Kiểm tra bộ chọn Mall trên thanh trên cùng. Dữ liệu và thống kê sẽ thay đổi theo Mall được chọn; Super Admin có thể chọn Tất cả Mall.')
add_step(4, 'Kiểm tra vai trò', 'Tên và vai trò xuất hiện ở góc phải. Nếu thiếu menu, liên hệ Super Admin để kiểm tra role và quyền Mall.')
doc.add_heading('Nguyên tắc an toàn', level=2)
for x in ['Không ghi nhận thanh toán nếu chưa có chứng từ và xác minh của Finance.', 'Không chạy seed/setup trên cơ sở dữ liệu cần giữ dữ liệu.', 'Kiểm tra Mall và đối tượng trước khi tạo hoặc phê duyệt hồ sơ.', 'Tải bằng chứng lên ticket/Fitout trước khi hoàn thành công việc.']: add_bullet(x)

doc.add_page_break(); doc.add_heading('2. Quy trình ERP end-to-end', level=1)
doc.add_paragraph('Luồng chuẩn: Lead CRM → Booking giữ mặt bằng → Proposal → Approval → Contract → Fitout → Tenant vận hành → Billing/SAP → Ticket & bảo trì.')
for i, (name, owner, done) in enumerate([
('Lead CRM','Leasing','Đủ thông tin liên hệ, nhu cầu, nguồn lead'),('Booking','Leasing','Unit được giữ đúng thời hạn'),('Proposal','Leasing/Manager','Điều khoản và giá thuê hoàn chỉnh'),('Approval','Cấp duyệt','Đủ lịch sử người duyệt và thời gian'),('Contract','Legal/Leasing','Hợp đồng hiệu lực và gắn tenant/unit'),('Fitout','Operation/Tenant','Checklist, hồ sơ, bằng chứng đầy đủ'),('Billing/SAP','Finance','Hóa đơn và đối soát hợp lệ'),('Ticket','Operation','Xử lý, bằng chứng, xác nhận đóng')],1):
    add_step(i, name, f'Chủ trì: {owner}. Điểm hoàn thành: {done}.')

add_module('3. Dashboard điều hành', '01-dashboard.png', 'Màn hình dành cho việc ra quyết định nhanh: sức khỏe danh mục, lấp đầy, công nợ, phê duyệt và ticket.', [('Chọn phạm vi Mall','Chọn một Mall hoặc Tất cả Mall ở thanh trên cùng.'),('Đọc vùng cần chú ý','Ưu tiên thẻ màu cảnh báo và số lượng cần xử lý.'),('Mở danh sách chi tiết','Nhấp thẻ KPI hoặc hành động để đi tới module nguồn.'),('Làm mới','Chọn Làm mới khi cần lấy dữ liệu tức thời.')])
add_module('4. Mặt bằng (Spaces)', '02-spaces.png', 'Tra cứu tồn kho, trạng thái, giá và sơ đồ mặt bằng.', [('Lọc dữ liệu','Chọn Mall, tầng, trạng thái hoặc ngành hàng.'),('Mở unit','Chọn một mặt bằng để xem diện tích, giá, tenant và lịch sử.'),('Cập nhật','Chỉ người có quyền mới được sửa cấu trúc hoặc giá.'),('Đối chiếu','Trước booking, xác nhận unit đang VACANT và không bị giữ bởi giao dịch khác.')])
add_module('5. CRM và Lead', '03-crm.png', 'Quản lý khách hàng tiềm năng từ tiếp nhận tới đủ điều kiện tạo booking.', [('Tạo lead','Nhập thương hiệu, liên hệ, nhu cầu diện tích và ngành hàng.'),('Phân công','Giao đúng nhân viên phụ trách và thời hạn hành động tiếp theo.'),('Cập nhật pipeline','Di chuyển trạng thái sau mỗi tương tác, ghi chú đầy đủ.'),('Chuyển booking','Chỉ chuyển khi nhu cầu và mặt bằng mục tiêu đã rõ.')])
add_module('6. Booking', '04-bookings.png', 'Giữ mặt bằng có thời hạn cho lead đủ điều kiện.', [('Tạo booking','Chọn lead và unit trống.'),('Kiểm tra xung đột','Đối chiếu trạng thái unit và các booking đang giữ.'),('Kích hoạt','Xác nhận thời hạn giữ và người phụ trách.'),('Chuyển đề xuất','Tạo Proposal từ booking để giữ liên kết dữ liệu.')])
add_module('7. Proposal', '05-proposals.png', 'Soạn, lưu, in và gửi đề xuất thuê.', [('Chọn hồ sơ','Mở proposal từ booking/lead liên quan.'),('Hoàn thiện điều khoản','Nhập giá, phí, thời hạn, ưu đãi và lịch thanh toán.'),('Lưu nháp','Lưu trước khi chuyển tab hoặc đóng popup.'),('In đề xuất','Dùng nút In/Xuất PDF; kiểm tra bản in không còn nút hoặc nội dung tương tác.'),('Gửi phê duyệt','Chỉ gửi khi form đầy đủ và số liệu đã được kiểm tra.')])
add_module('8. Approval', '06-approvals.png', 'Theo dõi hàng đợi duyệt, chi tiết hồ sơ và audit trail.', [('Mở hàng đợi','Lọc theo Chờ tôi duyệt, trạng thái và Mall.'),('Xem chi tiết','Kiểm tra Proposal, giá, điều khoản, người tạo và file liên quan.'),('Ra quyết định','Chọn Duyệt hoặc Từ chối; nhập lý do rõ ràng.'),('Kiểm tra lịch sử','Xác nhận ai đã duyệt, thời gian và từng cấp duyệt.'),('In form','Sau khi hoàn tất toàn bộ cấp, in form phê duyệt để lưu hồ sơ.')])
add_module('9. Contract', '07-contracts.png', 'Quản lý hợp đồng sau phê duyệt.', [('Tạo hợp đồng','Khởi tạo từ proposal đã duyệt để kế thừa dữ liệu.'),('Kiểm tra liên kết','Xác nhận tenant, unit, ngày hiệu lực và tiền thuê.'),('Đính kèm','Tải bản ký và tài liệu pháp lý.'),('Kích hoạt','Chỉ kích hoạt khi đủ phê duyệt và chữ ký.'),('Theo dõi hết hạn','Dùng bộ lọc và cảnh báo để chuẩn bị gia hạn.')])
add_module('10. Fitout', '08-fitout.png', 'Điều phối dự án thi công từ hồ sơ thiết kế tới khai trương.', [('Mở dự án','Chọn dự án để xem việc tiếp theo và tiến độ giai đoạn.'),('Phân công','Giao người phụ trách và mốc thời gian.'),('Hoàn thành checklist','Đánh dấu từng mục, đính kèm hồ sơ/bằng chứng bắt buộc.'),('Chuyển giai đoạn','Chỉ chuyển khi điều kiện và phê duyệt của giai đoạn hiện tại đạt.'),('Theo dõi rủi ro','Ưu tiên dự án trễ hạn, thiếu hồ sơ hoặc chưa phân công.')])
add_module('11. Ticket và bảo trì', '09-tickets.png', 'Tiếp nhận yêu cầu, quản lý SLA và kế hoạch bảo trì định kỳ.', [('Tạo yêu cầu','Chọn tenant/unit, loại, mức ưu tiên và mô tả.'),('Phân công','Giao người xử lý; kiểm tra hạn SLA.'),('Cập nhật','Ghi bình luận, trạng thái và ảnh bằng chứng.'),('Hoàn thành','Chỉ hoàn thành sau khi có kết quả và evidence.'),('Bảo trì định kỳ','Tạo kế hoạch, lịch nhắc, checklist, người chịu trách nhiệm và chu kỳ lặp.')])
add_module('12. Billing & AR', '10-billing.png', 'Theo dõi hóa đơn, công nợ, thu tiền và đối soát.', [('Lọc hóa đơn','Chọn Mall, tenant, trạng thái và kỳ.'),('Kiểm tra số dư','Đối chiếu tổng tiền, đã trả, balance và hạn thanh toán.'),('Ghi nhận thanh toán','Finance xác minh chứng từ trước khi ghi nhận.'),('Theo dõi quá hạn','Ưu tiên OVERDUE và PARTIALLY_PAID còn số dư.'),('Đối soát SAP','Kiểm tra trạng thái đồng bộ và lỗi tích hợp.')], ['Trong phiên bản QC này, tenant có thể gọi luồng ghi nhận thanh toán trực tiếp. Tạm thời yêu cầu Finance kiểm soát và xử lý bản vá Critical trước production.'])
add_module('13. Announcements', '11-announcements.png', 'Trung tâm truyền thông Mall cho cảnh báo, bảo trì, chính sách và sự kiện.', [('Chọn Mall','Nhân viên phải chọn Mall trước khi tạo thông báo.'),('Tạo thông báo','Nhập tiêu đề, nội dung, danh mục, ưu tiên, thời gian đăng/hết hạn.'),('Kiểm tra hiển thị','Tenant chỉ nhận thông báo thuộc Mall đang thuê.'),('Xem chi tiết','Chọn Xem chi tiết để đọc toàn bộ nội dung.'),('Xóa','Xác nhận trước khi xóa; lịch sử thao tác cần được quản trị theo quy định.')])
add_module('14. Tenant Portal', '13-tenant-portal.png', 'Không gian tự phục vụ dành cho khách thuê.', [('Xem tổng quan','Kiểm tra hợp đồng hiệu lực, công nợ, ticket và Fitout.'),('Xem hóa đơn','Lọc theo trạng thái và kiểm tra hạn thanh toán.'),('Gửi hỗ trợ','Chọn mặt bằng của chính tenant, loại yêu cầu, ưu tiên và mô tả.'),('Theo dõi xử lý','Mở ticket để bình luận, xem ảnh và trạng thái.'),('Theo dõi Fitout','Kiểm tra tiến độ, hồ sơ và việc cần hoàn thành.')])
add_module('15. Admin', '12-admin.png', 'ERP Control Center dành cho Super Admin.', [('Tài khoản','Tạo, khóa, đặt lại thông tin và gán role.'),('Quyền Mall','Giới hạn dữ liệu theo Mall cho từng người dùng.'),('Cấu trúc','Quản lý Mall, block, tầng, ngành hàng và giá.'),('Phân quyền','Kiểm tra role và quyền module trước khi cấp.'),('Approval Policy','Cấu hình cấp duyệt, thứ tự và ngưỡng.'),('Hệ thống','Quản lý branding và cấu hình vận hành.')], ['Mọi thay đổi Admin có thể ảnh hưởng toàn hệ thống; ghi nhận người thay đổi, thời gian và lý do.'])

doc.add_page_break(); doc.add_heading('16. Checklist vận hành hằng ngày', level=1)
for x in ['Kiểm tra health/readiness và trạng thái container.', 'Kiểm tra Approval, công nợ quá hạn, ticket SLA và Fitout trễ.', 'Kiểm tra log lỗi 5xx, job đồng bộ SAP và notification.', 'Xác nhận backup mới nhất và dung lượng lưu trữ.', 'Không chạy migrate/setup có seed trên DB chứa dữ liệu cần giữ.']: add_bullet(x)
doc.add_heading('Xử lý sự cố nhanh', level=2)
for title, detail in [('Không thấy menu','Kiểm tra role, permission và Mall access.'),('Dropdown không chọn được','Tải lại trang; kiểm tra dữ liệu nguồn và quyền API.'),('Số badge không khớp','Làm mới dữ liệu, bỏ filter và kiểm tra endpoint thống kê.'),('In bị lẫn nút','Dùng chức năng In/Xuất PDF của module, không in trực tiếp popup thao tác.'),('API 403','Xác nhận vai trò và phạm vi Mall; không cố dùng tài khoản khác để vượt quyền.')]: add_step('', title, detail)

doc.add_page_break(); doc.add_heading('17. Phụ lục kết quả QC ngày 19/07/2026', level=1)
doc.add_paragraph('Kết quả tự động: backend 44/44 test suites, 198/198 tests đạt; frontend 13/13 files, 61/61 tests đạt. UAT health/readiness/frontend HTTP 200 và 18 migration đồng bộ.')
doc.add_heading('Lỗi cần xử lý trước Production', level=2)
issues = [
('Critical','Tenant có thể ghi nhận payment trực tiếp và làm đổi trạng thái hóa đơn.'),
('Critical','migrate-uat bật SEED_DATABASE=true; seed có thao tác xóa dữ liệu nghiệp vụ.'),
('Critical','Tenant truy cập được API cấu hình Ticket SLA; payload sai gây 500 thay vì 403/400.'),
('High','IDOR ở escalation/rating ticket và thiếu kiểm tra ownership khi rate.'),
('High','Tenant xem maintenance toàn Mall; route thay đổi bảo trì thiếu role restriction.'),
('High','Staff được vào Tenant Portal nhưng một số API Billing trả 403 và CTA ticket không phù hợp.'),
('High','Container không có restart policy; PostgreSQL/Redis bind 0.0.0.0, Redis không auth.'),
]
t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.rows[0].cells[0].text, t.rows[0].cells[1].text = 'Mức độ', 'Phát hiện'
shade(t.rows[0].cells[0], LIGHT); shade(t.rows[0].cells[1], LIGHT)
for sev, issue in issues:
    c=t.add_row().cells; c[0].text=sev; c[1].text=issue
doc.add_paragraph('Báo cáo chi tiết đi kèm: QC_UI_REPORT.md, QC_BACKEND_REPORT.md và OPS_QC_REPORT.md.')

doc.core_properties.title = 'THISO Leasing Platform - Hướng dẫn sử dụng'
doc.core_properties.subject = 'Sổ tay vận hành ERP và kết quả QC'
doc.core_properties.author = 'THISO Leasing Platform QC & Operations'
doc.save(OUT)
print(OUT)
